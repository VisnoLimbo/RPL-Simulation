# -*- coding: utf-8 -*-
"""Replace Figure 1 with the hierarchical tree; update its caption and the
paragraph that described the previous geographic figure."""
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

SRC = "Customized_Routing_Simulator_Report.docx"
IMG = "results/dodag_hierarchical.png"

NEW_CAPTION = (
    "Figure 1. The DODAG formed in the baseline 20-node scenario, drawn as a "
    "hierarchical tree. The root occupies the top tier; each tier downward "
    "represents one further hop from the root and a rank higher by 768 (OF0). "
    "Arrows run from each node to its preferred parent — the direction of "
    "upward routing toward the root."
)
NEW_BODY = (
    "Figure 1 shows the DODAG formed in the baseline 20-node scenario, drawn "
    "as a hierarchical tree in which the root occupies the top tier and each "
    "tier downward corresponds to one further hop from the root. The twenty "
    "nodes are placed uniformly at random within the deployment area — "
    "with the root at its centre — producing a radio graph of 50 "
    "bidirectional links over which the DODAG is then constructed. The "
    "resulting DODAG has three tiers below the root: seven nodes attach "
    "directly to the root and obtain rank 1024, eight nodes attach at the "
    "second tier with rank 1792, and the four remaining nodes occupy the third "
    "tier with rank 2560. These rank values follow directly from OF0, under "
    "which the root holds rank 256 and each hop adds a fixed increment of 768."
)

doc = Document(SRC)

# ── 1. replace the Figure 1 image (first drawing in the document) ──────────
fig_para = None
for p in doc.paragraphs:
    if p._p.findall(".//" + qn("w:drawing")):
        fig_para = p
        break
for r in list(fig_para.runs):
    r._element.getparent().remove(r._element)
fig_para.add_run().add_picture(IMG, width=Inches(6.5))
print("Figure 1 image replaced")

# ── 2/3. update the caption and the descriptive body paragraph ─────────────
cap_done = body_done = False
for p in doc.paragraphs:
    t = p.text
    if not cap_done and t.startswith("Figure 1. "):
        p.runs[0].text = NEW_CAPTION
        for extra in list(p.runs)[1:]:
            extra._element.getparent().remove(extra._element)
        cap_done = True
    elif not body_done and t.startswith("Figure 1 shows"):
        p.runs[0].text = NEW_BODY
        for extra in list(p.runs)[1:]:
            extra._element.getparent().remove(extra._element)
        body_done = True
print(f"caption updated: {cap_done}   body paragraph updated: {body_done}")

doc.save(SRC)

# ── verify ─────────────────────────────────────────────────────────────────
d = Document(SRC)
print(f"verify: {len(d.inline_shapes)} images, {len(d.tables)} tables, "
      f"{len(d.paragraphs)} paragraphs")
s = d.inline_shapes[0]
print(f"Figure 1 display size: {s.width / 914400:.2f} x "
      f"{s.height / 914400:.2f} in")
